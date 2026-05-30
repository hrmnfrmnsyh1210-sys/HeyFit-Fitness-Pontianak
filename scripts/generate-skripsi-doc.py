"""Generate BAB II & BAB III skripsi document for Heyfit-Fitness."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt


OUTPUT = Path(__file__).resolve().parent.parent / "BAB_II_III_Skripsi_Heyfit.docx"

doc = Document()

# Page setup: A4, margins khas skripsi (kiri 4cm, atas/kanan/bawah 3cm)
section = doc.sections[0]
section.page_height = Cm(29.7)
section.page_width = Cm(21.0)
section.top_margin = Cm(3)
section.bottom_margin = Cm(3)
section.left_margin = Cm(4)
section.right_margin = Cm(3)

# Default Normal style: Times New Roman 12, line spacing 1.5
normal = doc.styles["Normal"]
normal.font.name = "Times New Roman"
normal.font.size = Pt(12)
rpr = normal.element.get_or_add_rPr()
rfonts = rpr.find(qn("w:rFonts"))
if rfonts is None:
    rfonts = OxmlElement("w:rFonts")
    rpr.append(rfonts)
for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
    rfonts.set(qn(attr), "Times New Roman")
normal.paragraph_format.line_spacing = 1.5
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(0)


def style_run(run, *, bold=False, italic=False, size=12):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    r = run._element
    rpr_ = r.get_or_add_rPr()
    rf = rpr_.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr_.append(rf)
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(a), "Times New Roman")


def add_heading(text, level=1, center=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    size = 14 if level == 1 else 12
    run = p.add_run(text)
    style_run(run, bold=True, size=size)
    return p


def add_paragraph(text, *, indent=True, justify=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if justify else WD_ALIGN_PARAGRAPH.LEFT
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    style_run(run)
    return p


def add_subheading(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    style_run(run, bold=True)
    return p


def add_numbered_item(number, bold_part, rest):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(f"{number} ")
    style_run(r1)
    r2 = p.add_run(bold_part)
    style_run(r2, bold=True)
    r3 = p.add_run(rest)
    style_run(r3)
    return p


def add_page_break():
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# ==================== BAB II ====================
add_heading("BAB II", level=1, center=True)
add_heading("LANDASAN TEORI", level=1, center=True)

add_heading("2.1 Konsep Dasar", level=2)

add_subheading("C. Program")

add_subheading("Definisi Program")
add_paragraph(
    "Program adalah serangkaian instruksi atau perintah yang ditulis dengan aturan tertentu "
    "dalam suatu bahasa pemrograman untuk memerintahkan komputer menyelesaikan tugas atau "
    "memecahkan suatu masalah. Program merupakan hasil dari proses pemrograman, yaitu "
    "kegiatan menulis, menguji, memperbaiki, dan memelihara kode yang membangun sebuah "
    "perangkat lunak."
)

add_subheading("Karakteristik Pembuatan Program")
add_paragraph(
    "Dalam membangun sebuah program yang baik, terdapat beberapa karakteristik yang harus "
    "diperhatikan:"
)
add_numbered_item("1.", "Kebenaran (Correctness)",
    " — program harus menghasilkan output yang sesuai dengan spesifikasi kebutuhan.")
add_numbered_item("2.", "Keandalan (Reliability)",
    " — program mampu berjalan stabil tanpa terjadinya kesalahan (error/bug) yang fatal.")
add_numbered_item("3.", "Efisiensi (Efficiency)",
    " — program dapat dijalankan dengan penggunaan sumber daya (memori, processor, jaringan) yang optimal.")
add_numbered_item("4.", "Kemudahan Penggunaan (Usability)",
    " — antarmuka program mudah dipahami dan digunakan oleh pengguna akhir.")
add_numbered_item("5.", "Kemudahan Pemeliharaan (Maintainability)",
    " — kode program terstruktur sehingga mudah dimodifikasi dan dikembangkan di kemudian hari.")
add_numbered_item("6.", "Portabilitas (Portability)",
    " — program dapat dijalankan pada berbagai platform/perangkat (responsif terhadap desktop, tablet, maupun smartphone).")

add_subheading("Bahasa Pemrograman yang Digunakan")
add_paragraph(
    "Pada penelitian ini, penulis menggunakan beberapa bahasa pemrograman dan teknologi "
    "pendukung untuk membangun website Heyfit-Fitness, antara lain:"
)
add_numbered_item("1.", "TypeScript",
    " — superset dari JavaScript yang menambahkan fitur static typing. Digunakan sebagai bahasa utama pengembangan karena mendukung tipe data yang ketat (strict mode) sehingga mengurangi potensi kesalahan pada saat kompilasi.")
add_numbered_item("2.", "JavaScript",
    " — bahasa pemrograman berbasis script yang dijalankan pada sisi client (browser) untuk membuat halaman web menjadi interaktif dan dinamis.")
add_numbered_item("3.", "HTML (HyperText Markup Language)",
    " — bahasa markup standar untuk menyusun struktur halaman web.")
add_numbered_item("4.", "CSS (Cascading Style Sheets)",
    " — bahasa yang digunakan untuk mengatur tampilan dan tata letak elemen HTML. Pada penelitian ini, penulisan CSS dibantu oleh framework Tailwind CSS dengan pendekatan utility-first.")
add_numbered_item("5.", "SQL (Structured Query Language)",
    " — bahasa query yang digunakan untuk mengelola data pada basis data relasional MySQL.")
add_paragraph(
    "Sebagai framework, penulis menggunakan Nuxt 3 yang berbasis Vue.js 3 untuk membangun "
    "antarmuka pengguna (frontend) sekaligus server-side (backend) dalam satu basis kode "
    "(full-stack), dengan dukungan Drizzle ORM untuk akses basis data."
)

add_subheading("D. Basis Data")

add_subheading("Definisi Basis Data")
add_paragraph(
    "Basis data (database) adalah kumpulan data yang saling berhubungan dan terorganisasi "
    "sedemikian rupa sehingga dapat disimpan, dikelola, dan diakses kembali dengan mudah "
    "dan cepat. Basis data berfungsi untuk menghindari duplikasi data, menjaga konsistensi "
    "data, serta mempermudah pencarian dan manipulasi data oleh aplikasi."
)

add_subheading("Aplikasi Basis Data yang Digunakan")
add_paragraph(
    "Aplikasi basis data yang digunakan pada website Heyfit-Fitness adalah MySQL (melalui "
    "layanan kompatibel TiDB). MySQL merupakan salah satu Relational Database Management "
    "System (RDBMS) bersifat open source yang menggunakan bahasa SQL untuk mengelola data. "
    "Penulis menggunakan MySQL karena memiliki performa yang baik, dukungan komunitas yang "
    "luas, serta kompatibel dengan berbagai bahasa pemrograman."
)
add_paragraph(
    "Untuk mempermudah pengelolaan skema dan query, penulis menggunakan Drizzle ORM sebagai "
    "Object Relational Mapping yang memetakan tabel basis data ke dalam objek TypeScript, "
    "dilengkapi dengan Drizzle Kit untuk migrasi skema basis data."
)

add_subheading("E. Model Pengembangan Perangkat Lunak")
add_paragraph(
    "Model pengembangan perangkat lunak yang digunakan dalam penelitian ini adalah model "
    "Waterfall (air terjun). Model waterfall merupakan salah satu model klasik dalam "
    "Software Development Life Cycle (SDLC) yang bersifat sistematis dan berurutan, di mana "
    "setiap tahapan harus diselesaikan terlebih dahulu sebelum berlanjut ke tahapan "
    "berikutnya. Model ini terdiri dari lima tahapan utama, yaitu analisis kebutuhan, "
    "desain, pembuatan kode program (implementasi), pengujian, serta pendukung (support) "
    "atau pemeliharaan (maintenance)."
)
add_paragraph(
    "Alasan penulis memilih model waterfall adalah karena kebutuhan dari website "
    "Heyfit-Fitness telah didefinisikan secara jelas sejak awal, ruang lingkup terbatas, "
    "dan tidak terdapat perubahan kebutuhan yang signifikan selama proses pengembangan."
)

add_heading("2.2 Teori Pendukung", level=2)

add_subheading("A. Entity Relationship Diagram (ERD)")

add_subheading("Definisi ERD")
add_paragraph(
    "Entity Relationship Diagram (ERD) adalah suatu diagram yang digunakan untuk memodelkan "
    "struktur data dan hubungan antar data dalam sebuah basis data secara konseptual. ERD "
    "membantu perancang sistem dalam memahami entitas-entitas yang terlibat beserta "
    "hubungan (relasi) di antara entitas tersebut sebelum diimplementasikan ke dalam basis "
    "data fisik."
)

add_subheading("Komponen ERD")
add_paragraph("ERD memiliki beberapa komponen utama, yaitu:")
add_numbered_item("1.", "Entitas (Entity)",
    " — objek nyata atau abstrak yang datanya akan disimpan, digambarkan dengan bentuk persegi panjang. Contoh: User, Member, Kelas, Pembayaran.")
add_numbered_item("2.", "Atribut (Attribute)",
    " — karakteristik atau properti yang melekat pada suatu entitas, digambarkan dengan bentuk elips. Contoh atribut pada entitas Member: id_member, nama, email, no_telp.")
add_numbered_item("3.", "Relasi (Relationship)",
    " — hubungan antara dua atau lebih entitas, digambarkan dengan bentuk belah ketupat. Contoh: Member mendaftar Kelas.")
add_numbered_item("4.", "Kardinalitas (Cardinality)",
    " — menggambarkan jumlah keterhubungan antar entitas, yaitu one-to-one (1:1), one-to-many (1:M), dan many-to-many (M:N).")

add_subheading("Logical Record Structure (LRS)")
add_paragraph(
    "Logical Record Structure (LRS) adalah representasi dari struktur record-record pada "
    "tabel-tabel yang terbentuk dari hasil relasi antar entitas pada ERD. LRS menggambarkan "
    "tipe record yang ada, termasuk primary key dan foreign key yang menghubungkan antar "
    "tabel, sehingga lebih dekat dengan bentuk implementasi fisik basis data."
)

add_subheading("B. Unified Modelling Language (UML)")

add_subheading("Definisi UML")
add_paragraph(
    "Unified Modelling Language (UML) adalah sebuah bahasa pemodelan visual yang digunakan "
    "untuk menspesifikasikan, memvisualisasikan, membangun, dan mendokumentasikan rancangan "
    "dari sebuah sistem perangkat lunak berorientasi objek. UML menyediakan berbagai jenis "
    "diagram yang dapat digunakan sesuai kebutuhan analisis maupun perancangan."
)

add_subheading("Use Case Diagram")
add_paragraph(
    "Use Case Diagram adalah diagram yang menggambarkan interaksi antara aktor (pengguna "
    "sistem) dengan fungsi-fungsi yang disediakan oleh sistem. Pada website Heyfit-Fitness, "
    "terdapat tiga aktor utama yaitu Member, Admin, dan Owner, dengan use case seperti "
    "login, pendaftaran keanggotaan, pendaftaran kelas, pemindaian QR untuk absensi, "
    "manajemen pembayaran, dan lain sebagainya."
)

add_subheading("Activity Diagram")
add_paragraph(
    "Activity Diagram adalah diagram yang menggambarkan alur kerja (workflow) atau "
    "aktivitas dari sebuah sistem atau proses bisnis dari awal hingga akhir. Diagram ini "
    "menampilkan urutan aktivitas, percabangan kondisi (decision), dan paralelisme "
    "aktivitas."
)

add_subheading("Class Diagram")
add_paragraph(
    "Class Diagram adalah diagram yang menggambarkan struktur statis dari sebuah sistem "
    "berbasis objek, meliputi kelas-kelas yang ada, atribut, metode, serta hubungan antar "
    "kelas. Class diagram digunakan sebagai dasar dalam merancang struktur data dan "
    "komponen aplikasi."
)

add_subheading("Sequence Diagram")
add_paragraph(
    "Sequence Diagram adalah diagram yang menggambarkan interaksi antar objek dalam sebuah "
    "sistem yang disusun berdasarkan urutan waktu (kronologis). Diagram ini menjelaskan "
    "pesan (message) yang dikirim dan diterima antar objek pada saat sebuah skenario atau "
    "use case dijalankan."
)

add_page_break()

# ==================== BAB III ====================
add_heading("BAB III", level=1, center=True)
add_heading("METODE PENELITIAN", level=1, center=True)

add_heading("Model Pengembangan Perangkat Lunak (Waterfall)", level=2)

add_paragraph(
    "Dalam pengembangan website Heyfit-Fitness, penulis menggunakan model pengembangan "
    "perangkat lunak Waterfall yang terdiri dari beberapa tahapan sebagai berikut:"
)

add_subheading("1. Analisis Kebutuhan Perangkat Lunak")
add_paragraph(
    "Dalam tahap awal, penulis melakukan analisa terhadap kebutuhan dalam membangun "
    "website Heyfit-Fitness, yaitu sebuah platform manajemen pusat kebugaran (gym/fitness) "
    "berbasis web. Pengguna yang terlibat terdiri dari tiga kategori, yaitu Member, Admin, "
    "dan Owner."
)
add_numbered_item("a.", "Member",
    " adalah pengguna yang telah memiliki akun terdaftar pada website. Member dapat melakukan pendaftaran keanggotaan, perpanjangan keanggotaan, melihat daftar kelas (Yoga, Pilates, dan lain-lain), mendaftar kelas, melihat informasi fasilitas, melakukan pembayaran, serta melakukan absensi melalui scan kode QR.")
add_numbered_item("b.", "Admin",
    " adalah pengguna yang bertugas mengelola data master pada sistem, meliputi data member, data instruktur, data kelas, data pembayaran, serta memverifikasi absensi member dengan melakukan scan kode QR. Admin dapat menambah, mengubah, dan menghapus data tersebut.")
add_numbered_item("c.", "Owner",
    " adalah pengguna dengan hak akses tertinggi yang bertugas memantau keseluruhan operasional dan performa bisnis melalui dashboard khusus owner.")

add_subheading("2. Desain")
add_paragraph(
    "Setelah tahap analisa kebutuhan terpenuhi, dalam perancangan ini dilakukan proses "
    "perancangan sistem dan perangkat lunak dengan membagi kebutuhan perangkat keras "
    "(hardware) dan perangkat lunak (software). Proses ini berfokus pada rancangan antar "
    "muka, struktur basis data, dan rancangan struktur navigasi."
)
add_paragraph(
    "Rancangan antar muka meliputi tampilan untuk member (halaman beranda, keanggotaan, "
    "kelas, fasilitas), tampilan admin (dashboard admin, manajemen member, instruktur, "
    "kelas, pembayaran, scan absensi), serta tampilan owner. Sedangkan rancangan struktur "
    "basis data meliputi Entity Relationship Diagram (ERD), Logical Record Structure (LRS), "
    "dan Spesifikasi File. Pemodelan sistem digambarkan menggunakan Unified Modelling "
    "Language (UML) yang terdiri dari Use Case Diagram, Activity Diagram, Class Diagram, "
    "dan Sequence Diagram."
)

add_subheading("3. Pembuatan Kode Program")
add_paragraph(
    "Setelah selesai merancang desain, penulis melakukan pengkodingan pada program yang "
    "dirancang. Adapun teknologi yang digunakan adalah HTML, CSS dengan framework Tailwind "
    "CSS, TypeScript dan JavaScript sebagai bahasa pemrograman, Nuxt 3 (berbasis Vue.js 3) "
    "sebagai framework full-stack, Drizzle ORM untuk pemetaan basis data, MySQL (TiDB) "
    "sebagai Database Management System, serta Visual Studio Code sebagai editor untuk "
    "menulis kode program. Untuk fitur absensi digunakan library html5-qrcode dan qrcode "
    "untuk pembuatan serta pembacaan kode QR."
)

add_subheading("4. Pengujian")
add_paragraph(
    "Kemudian untuk memastikan program yang pengkodingannya telah selesai berjalan dengan "
    "baik, penulis melakukan pengujian dengan menggunakan metode Blackbox Testing. Metode "
    "ini digunakan untuk menguji fungsionalitas sistem tanpa melihat struktur kode di "
    "dalamnya, dengan cara mencocokkan input yang diberikan dengan output yang dihasilkan."
)
add_paragraph(
    "Pengujian dilakukan terhadap hak akses Member (pendaftaran akun, login, pendaftaran "
    "keanggotaan, perpanjangan keanggotaan, pendaftaran kelas, pembayaran, absensi melalui "
    "scan QR), hak akses Admin (login admin, manajemen data member, instruktur, kelas, "
    "pembayaran, scan QR absensi), serta hak akses Owner (login owner dan akses dashboard "
    "owner). Tujuan pengujian ini adalah untuk menemukan kesalahan (error) yang mungkin "
    "terjadi agar dapat diperbaiki sebelum program digunakan oleh pengguna akhir."
)

add_subheading("5. Pendukung (Support) atau Pemeliharaan (Maintenance)")
add_paragraph(
    "Pada tahap ini dilakukan dukungan terhadap program pada lingkungan operasionalnya "
    "serta proses pemeliharaan. Pemeliharaan mencakup koreksi dari berbagai error yang "
    "tidak ditemukan pada tahapan sebelumnya, melakukan perbaikan atas implementasi unit "
    "sistem, pengembangan layanan sistem, serta penambahan persyaratan-persyaratan baru "
    "sesuai kebutuhan pengguna di masa mendatang, seperti penambahan fitur baru pada "
    "manajemen kelas, paket keanggotaan, maupun integrasi metode pembayaran lainnya."
)

doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
